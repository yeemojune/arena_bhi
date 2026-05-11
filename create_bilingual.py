#!/usr/bin/env python3
"""
Create a bilingual (English-Chinese) version of the BHI NeurIPS 2026 paper.
Format: English text followed by Chinese translation.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

def add_bilingual_para(doc, en_text, cn_text, style=None, en_font_size=10.5, cn_font_size=10.5, en_bold=False, cn_bold=False, alignment=None):
    """Add a bilingual paragraph: English first, then Chinese below."""
    # English paragraph
    p_en = doc.add_paragraph()
    if alignment:
        p_en.alignment = alignment
    run_en = p_en.add_run(en_text)
    run_en.font.size = Pt(en_font_size)
    run_en.font.bold = en_bold
    run_en.font.color.rgb = RGBColor(0, 0, 0)
    run_en.font.name = 'Times New Roman'
    
    # Chinese paragraph
    p_cn = doc.add_paragraph()
    if alignment:
        p_cn.alignment = alignment
    run_cn = p_cn.add_run(cn_text)
    run_cn.font.size = Pt(cn_font_size)
    run_cn.font.bold = cn_bold
    run_cn.font.color.rgb = RGBColor(0, 0, 128)
    run_cn.font.name = 'SimSun'

def add_heading_bilingual(doc, en_text, cn_text, level=1):
    """Add a bilingual heading."""
    h = doc.add_heading(level=level)
    run_en = h.add_run(en_text + "\n")
    run_en.font.size = Pt(14 if level == 1 else 12)
    run_en.font.name = 'Times New Roman'
    run_cn = h.add_run(cn_text)
    run_cn.font.size = Pt(13 if level == 1 else 11)
    run_cn.font.color.rgb = RGBColor(0, 0, 128)
    run_cn.font.name = 'SimHei'

def add_separator(doc):
    """Add a visual separator between sections."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

doc = Document()

# ============ TITLE ============
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("Benchmark Health Index: A Systematic Framework for Benchmarking the Benchmarks of LLMs")
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'

p_title_cn = doc.add_paragraph()
p_title_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title_cn.add_run("基准健康指数：一个系统性的 LLM 基准评测框架")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 0, 128)
run.font.name = 'SimHei'

doc.add_paragraph()

# Authors
add_bilingual_para(doc, "Anonymous Author(s)", "匿名作者", en_font_size=11, cn_font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_bilingual_para(doc, "Affiliation", "所属机构", en_font_size=11, cn_font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_separator(doc)

# ============ ABSTRACT ============
add_heading_bilingual(doc, "Abstract", "摘要", level=1)

add_bilingual_para(doc,
    "Large Language Models (LLMs) are advancing rapidly, yet the benchmarks used to measure this progress are becoming increasingly unreliable. Score inflation and selective reporting have eroded the authority of standard benchmarks, leaving the community uncertain about which evaluation results remain trustworthy. We introduce the Benchmark Health Index (BHI), a pure data-driven framework for auditing evaluation sets along three orthogonal and complementary axes: (1) Capability Discrimination, measuring how sharply a benchmark separates model performance beyond noise; (2) Anti-Saturation, estimating remaining headroom before ceiling effects erode resolution and thus the benchmark's expected longevity; and (3) Impact, measuring capability-weighted benchmark adoption across the model ecosystem.",
    "大语言模型（LLM）正在快速进步，然而用于衡量这一进展的基准评测正变得日益不可靠。分数膨胀和选择性报告已侵蚀了标准基准的权威性，使社区对哪些评估结果仍可信赖感到不确定。我们提出了基准健康指数（Benchmark Health Index, BHI），一个纯数据驱动的框架，沿三个正交且互补的维度对评估集进行审计：（1）能力区分度，衡量基准在噪声之上分离模型性能的敏锐程度；（2）抗饱和度，估计天花板效应侵蚀分辨率之前剩余的提升空间，即基准的预期寿命；（3）影响力，衡量跨模型生态系统的能力加权基准采用率。")

add_bilingual_para(doc,
    "By distilling 158 validated benchmarks from the technical reports of 117 representative models released between 2025 and April 2026, we systematically characterize the evaluation landscape. BHI is the first framework to quantify benchmark health at a macro level, providing a well-grounded quantitative basis for benchmark selection and supporting the dynamic filtering and continuous updating of evaluation sets.",
    "通过从 2025 年至 2026 年 4 月期间发布的 117 个代表性模型的技术报告中提炼出 158 个经过验证的基准，我们系统地刻画了评估生态的全貌。BHI 是首个在宏观层面量化基准健康状况的框架，为基准选择提供了充分量化的依据，并支持评估集的动态筛选和持续更新。")

add_bilingual_para(doc,
    "Our comprehensive analysis not only identifies benchmarks that remain highly credible but also exposes pervasive structural issues. By establishing a rigorous foundation for selection and management, BHI elevates evaluation practice from heuristic-based judgment to a quantifiable and interpretable scientific process.",
    "我们的综合分析不仅识别出仍具高度可信度的基准，还揭示了普遍存在的结构性问题。通过建立严谨的选择与管理基础，BHI 将评估实践从基于启发式的判断提升为可量化和可解释的科学过程。")

add_separator(doc)

# ============ 1 INTRODUCTION ============
add_heading_bilingual(doc, "1  Introduction", "1  引言", level=1)

add_bilingual_para(doc,
    "The progress of LLMs is inseparable from benchmarks. Benchmarks define the trajectory of scientific advancement for the research community, provide capability references for industry, and offer essential guidance for allocating and optimizing computational resources. Ideally, a benchmark should function as a precise measurement instrument, delivering stable, reproducible, and interpretable discriminative signals across capability levels. However, the evaluation ecosystem has exhibited systematic distortions in recent years.",
    "LLM 的进步与基准密不可分。基准为研究界定义了科学进步的方向，为产业界提供能力参考，并为分配和优化计算资源提供必要指导。理想情况下，基准应作为精确的测量工具，在不同能力水平上提供稳定、可复现和可解释的区分信号。然而，近年来评估生态系统已呈现出系统性扭曲。")

add_bilingual_para(doc,
    "Many canonical benchmarks are approaching a compressed performance ceiling: top-tier models cluster within a narrow score band, so small gaps are often driven by statistical noise rather than substantive architectural differences. Moreover, under selective reporting and marketing-oriented practices, citation frequency no longer reliably reflects technical rigor or the long-term evaluative utility of a benchmark. These trends call for a shift from model-centric comparison to benchmark-centric auditing, treating benchmarks as public infrastructure with a lifecycle and quantitatively testing their validity and degradation dynamics.",
    "许多经典基准正逼近压缩的性能天花板：顶级模型聚集在狭窄的分数区间内，因此微小的差距往往源于统计噪声而非实质性的架构差异。此外，在选择性报告和营销导向的实践下，引用频率不再可靠地反映基准的技术严谨性或长期评估效用。这些趋势呼吁从以模型为中心的比较转向以基准为中心的审计，将基准视为具有生命周期的公共基础设施，并定量检验其有效性和退化动态。")

add_bilingual_para(doc,
    "Current work has laid important foundations in holistic evaluation frameworks, dataset construction quality, and cross-benchmark consistency. Yet most existing approaches remain limited to static specifications or local consistency analyses, lacking an actionable, quantitative, and reproducible yardstick for horizontal comparison and full-lifecycle management of benchmarks at the ecosystem scale. The community thus needs a global metric system for benchmark health: one that not only quantifies a benchmark's current discriminative capability, but also evaluates its remaining effective and ecological impact, providing a quantifiable basis for benchmark selection.",
    "现有工作在整体评估框架、数据集构建质量和跨基准一致性方面奠定了重要基础。然而，大多数现有方法仍局限于静态规范或局部一致性分析，缺乏一个可操作、可量化、可复现的标尺，用于在生态系统层面进行基准的横向比较和全生命周期管理。因此，社区需要一个基准健康的全局度量体系：不仅量化基准当前的区分能力，还评估其剩余有效性和生态影响，为基准选择提供可量化的依据。")

add_bilingual_para(doc,
    "Formalizing benchmark health faces three technical challenges: (1) Low Discrimination & Score Compression: many benchmarks can no longer clearly separate strong models because their scores are compressed near the top; (2) Ceiling Effect & Lack of Headroom: as benchmark data is absorbed into training or repeatedly optimized for, the remaining challenge space shrinks, weakening long-term progress tracking; and (3) Popularity Bias & Adoption Decoupling: frequent citation or reporting does not necessarily mean that a benchmark is widely adopted by high-capability models or remains technically informative. Together, these issues lead the community to overuse outdated benchmarks and overlook more valuable ones.",
    "形式化基准健康面临三大技术挑战：（1）低区分度与分数压缩：许多基准不再能清晰分离强模型，因为其分数被压缩在顶部附近；（2）天花板效应与缺乏提升空间：随着基准数据被训练吸收或反复优化，剩余的挑战空间缩小，削弱了长期进展追踪；（3）流行度偏差与采用脱钩：频繁的引用或报告不一定意味着该基准被高能力模型广泛采用或仍具技术信息量。这些问题共同导致社区过度使用过时的基准，而忽视更有价值的基准。")

add_bilingual_para(doc,
    "To address this, we introduce Benchmark Health Index, a data-driven auditing framework that characterizes benchmark validity and usability across three dimensions: (1) Capability Discrimination, which measures whether a benchmark can clearly separate stronger models from weaker ones; (2) Anti-Saturation, which evaluates whether a benchmark still remains challenging for current frontier models and is unlikely to be exhausted soon; and (3) Impact, which measures whether a benchmark is widely adopted by high-capability models in real evaluation practice. By aggregating these, BHI identifies effective benchmarks, demotes saturated tests to basic sanity checks, and designates high-difficulty anchors for frontier capability tracking.",
    "为解决这一问题，我们提出了基准健康指数，一个数据驱动的审计框架，沿三个维度刻画基准的有效性和可用性：（1）能力区分度，衡量基准是否能清晰分离更强与更弱的模型；（2）抗饱和度，评估基准是否对当前前沿模型仍具挑战性且不太可能很快被耗尽；（3）影响力，衡量基准是否被高能力模型在实际评估实践中广泛采用。通过聚合这些维度，BHI 识别有效基准，将饱和的测试降级为基础健全性检查，并指定高难度锚点用于追踪前沿能力。")

add_bilingual_para(doc,
    "We conducted a systematic meta-analysis of 158 benchmarks extracted from the technical reports of 117 representative LLMs released between 2025 and April 2026. Our study distinguishes effective benchmarks that retain high discriminative power from those that, despite being widely used, have become obsolete. Furthermore, we provide a systematic examination of benchmark distribution and evolutionary trends across diverse task domains.",
    "我们对从 2025 年至 2026 年 4 月期间发布的 117 个代表性 LLM 技术报告中提取的 158 个基准进行了系统性元分析。我们的研究区分了保留高区分力的有效基准与尽管被广泛使用但已过时的基准。此外，我们还对跨多元任务领域的基准分布和演化趋势进行了系统性审视。")

add_bilingual_para(doc,
    "Our contributions are summarized as follows:\n• We propose the BHI framework, which quantifies benchmark effectiveness and longevity along three dimensions: Capability Discrimination, Anti-Saturation, and Impact.\n• We establish an objective, data-driven auditing protocol that implements a Leave-One-Benchmark-Out calibration strategy and automated weighting mechanisms to eliminate subjective bias and ensure methodological rigor.\n• Through large-scale meta-analysis and diagnostic case studies over 117 models and 158 benchmarks, we show that BHI separates saturated from frontier benchmarks, explains their degradation patterns, and offers actionable guidance for benchmark governance.",
    "我们的贡献总结如下：\n• 我们提出了 BHI 框架，沿三个维度量化基准的有效性和寿命：能力区分度、抗饱和度和影响力。\n• 我们建立了一个客观的、数据驱动的审计协议，实施了留一基准校准策略和自动化加权机制，以消除主观偏差并确保方法论的严谨性。\n• 通过对 117 个模型和 158 个基准的大规模元分析和诊断案例研究，我们展示了 BHI 能区分饱和基准与前沿基准，解释其退化模式，并为基准治理提供可操作的指导。")

add_separator(doc)

# ============ 2 RELATED WORK ============
add_heading_bilingual(doc, "2  Related Work", "2  相关工作", level=1)

add_bilingual_para(doc,
    "The rapid evolution of LLMs has catalyzed a paradigm shift in evaluation, transitioning from static, model-centric leaderboards to holistic benchmark suites for multi-dimensional capabilities. However, current benchmarks are plagued by severe score inflation and upper-tail compression. Despite ongoing efforts to mitigate benchmark degradation and data contamination, a unified diagnostic framework to monitor benchmark efficacy over time remains elusive. To address this gap, we propose a systematic framework to quantify benchmark capability, enabling high-resolution efficacy measurement throughout the dynamic co-evolution of benchmarks and models.",
    "LLM 的快速演化催化了评估范式的转变，从静态的、以模型为中心的排行榜转向多维能力的整体基准套件。然而，当前基准受到严重的分数膨胀和上尾压缩的困扰。尽管在缓解基准退化和数据污染方面持续努力，但一个用于随时间监测基准效力的统一诊断框架仍然难以实现。为填补这一空白，我们提出了一个系统性框架来量化基准能力，使在基准与模型动态协同演化的过程中能够进行高分辨率的效力测量。")

add_separator(doc)

# ============ 3 METHODOLOGY ============
add_heading_bilingual(doc, "3  Methodology", "3  方法论", level=1)

add_bilingual_para(doc,
    "The BHI is constructed through a data-driven aggregation of three distinct evaluation axes. In this section, we define the mathematical formulations and explain the statistical rationale for each metric.",
    "BHI 通过数据驱动聚合三个不同的评估维度来构建。在本节中，我们定义了数学公式并解释了每个度量的统计原理。")

# 3.1
add_heading_bilingual(doc, "3.1  Capability Discrimination (SDisc)", "3.1  能力区分度（SDisc）", level=2)

add_bilingual_para(doc,
    "SDisc quantifies a benchmark b ∈ B's resolution in distinguishing latent model performance through micro-resolution and macro-bandwidth. All model scores are pre-normalized to a [0,100] scale for consistency.",
    "SDisc 通过微观分辨率和宏观带宽来量化基准 b ∈ B 区分潜在模型性能的分辨率。所有模型分数预归一化至 [0,100] 尺度以保持一致性。")

# 3.1.1
add_heading_bilingual(doc, "3.1.1  Effective Differentiation Ratio (EDR)", "3.1.1  有效区分比（EDR）", level=3)

add_bilingual_para(doc,
    "EDR measures micro-resolution as the proportion of all possible pairwise combinations of models whose performance gap exceeds a noise-filtering threshold. For N models, it is defined as:\nEDR(b) = P_{1≤i<j≤N} I(|Score_i − Score_j| > δ) / (N(N − 1)/2)\nwhere δ = 0.02 · (max(Score) − min(Score)) is an adaptive threshold representing 2% of the observed score range. Sensitivity analysis confirms that SDisc is highly robust across δ ∈ [0.5%,5.0%], ensuring it captures intrinsic discriminative signals rather than measurement noise.",
    "EDR 通过计算所有可能模型对中性能差距超过噪声过滤阈值的比例来衡量微观分辨率。对于 N 个模型，其定义为：\nEDR(b) = P_{1≤i<j≤N} I(|Score_i − Score_j| > δ) / (N(N − 1)/2)\n其中 δ = 0.02 · (max(Score) − min(Score)) 是代表观察分数范围 2% 的自适应阈值。敏感性分析证实 SDisc 在 δ ∈ [0.5%,5.0%] 范围内高度稳健，确保其捕获的是内在区分信号而非测量噪声。")

# 3.1.2
add_heading_bilingual(doc, "3.1.2  Robust Coefficient of Variation (RCV)", "3.1.2  稳健变异系数（RCV）", level=3)

add_bilingual_para(doc,
    "RCV characterizes macro-bandwidth by measuring the effective performance spread using the middle 80% of the score distribution for the N models, thereby mitigating the influence of extreme outliers:\nRCV(b) = (P90(Score) − P10(Score)) / 100",
    "RCV 通过使用 N 个模型分数分布的中间 80% 来衡量有效性能扩展，从而减轻极端异常值的影响，以刻画宏观带宽：\nRCV(b) = (P90(Score) − P10(Score)) / 100")

# 3.1.3
add_heading_bilingual(doc, "3.1.3  Normalization", "3.1.3  归一化", level=3)

add_bilingual_para(doc,
    "Prior to final fusion, we apply min-max normalization to both EDR and RCV to map them onto a common interval [0,1]. To ensure the final score accurately reflects the informative contribution of each metric, we aggregate these indicators using Standard Deviation-based Weighting (SDM). This approach amplifies relative contrast while mitigating statistical variance shrinkage in benchmarks with sparse model coverage. The final score for benchmark b is computed as:\nSDisc(b) = Σ_{k∈{EDR,RCV}} w_k · Norm(k)\nDetailed mathematical derivations for the SDM weighting and the normalization process are provided in Appendix D.",
    "在最终融合之前，我们对 EDR 和 RCV 应用最小-最大归一化，将它们映射到公共区间 [0,1]。为确保最终分数准确反映每个指标的信息贡献，我们使用基于标准差的加权（SDM）来聚合这些指标。该方法放大了相对对比度，同时减轻了模型覆盖稀疏的基准中的统计方差收缩。基准 b 的最终分数计算为：\nSDisc(b) = Σ_{k∈{EDR,RCV}} w_k · Norm(k)\nSDM 加权和归一化过程的详细数学推导见附录 D。")

# 3.2
add_heading_bilingual(doc, "3.2  Model Capability Calibration", "3.2  模型能力校准", level=2)

add_bilingual_para(doc,
    "To provide a fair reference for Anti-Saturation (Section 3.3) and Impact (Section 3.4), we establish a system to calibrate model capabilities. To prevent self-referential bias caused by evaluating a benchmark using model strengths derived from the same data pool, we implement a Leave-One-Benchmark-Out (LOBO) strategy where capability scores θ_i for any target benchmark b ∈ B are calculated strictly using data from the remaining B \\ {b} benchmarks, thereby ensuring methodological independence through out-of-sample evaluation.",
    "为抗饱和度（第 3.3 节）和影响力（第 3.4 节）提供公平的参考，我们建立了一个模型能力校准系统。为防止使用同一数据池推导出的模型强度来评估基准时产生的自参照偏差，我们实施了留一基准（LOBO）策略，其中任何目标基准 b ∈ B 的能力分数 θ_i 严格使用剩余 B \\ {b} 基准的数据计算，从而通过样本外评估确保方法论的独立性。")

# 3.2.1
add_heading_bilingual(doc, "3.2.1  LOBO-adjusted Win Rate", "3.2.1  LOBO 调整胜率", level=3)

add_bilingual_para(doc,
    "We leverage a pairwise battle mechanism to ensure a robust and unified standardization of scores across the diverse array of metrics. For model i on benchmark b, its LOBO-adjusted Win Rate represents the average probability of defeating all opponents across the set of benchmarks excluding b, where Win_{i,b'} and Tie_{i,b'} denote the number of wins and ties of model i on benchmark b', and N_{b'} is the total number of models evaluated on b'. This win rate ensures consistency across heterogeneous evaluation scales.",
    "我们利用成对对战机制来确保跨多样指标的分数具有稳健统一的标准化。对于基准 b 上的模型 i，其 LOBO 调整胜率表示在排除 b 的基准集合上击败所有对手的平均概率，其中 Win_{i,b'} 和 Tie_{i,b'} 分别表示模型 i 在基准 b' 上的获胜和平局次数，N_{b'} 是在 b' 上评估的模型总数。该胜率确保了跨异构评估尺度的一致性。")

# 3.2.2
add_heading_bilingual(doc, "3.2.2  The Fourth-root Log-Balance Model", "3.2.2  四次根对数平衡模型", level=3)

add_bilingual_para(doc,
    "We integrate the win rate with the model's participation density to ensure that capability scores rigorously reflect both model performance and statistical reliability. For benchmark b, the calibrated capability of model i ∈ N is defined as:\nθ_i^{(-b)} = (WinRate_i^{(-b)})^{1/4} · (1 + log(n_i^{(-b)}) / log(N_B))\nwhere n_i^{(-b)} is the number of benchmarks model i has participated in within the LOBO reference pool B \\ {b}, and N_B is the total number of benchmarks in that reference pool (i.e., |B| − 1). The constant 1 is added to logarithmic terms for numerical stability.",
    "我们将胜率与模型的参与密度整合，确保能力分数严格反映模型性能和统计可靠性。对于基准 b，模型 i ∈ N 的校准能力定义为：\nθ_i^{(-b)} = (WinRate_i^{(-b)})^{1/4} · (1 + log(n_i^{(-b)}) / log(N_B))\n其中 n_i^{(-b)} 是模型 i 在 LOBO 参考池 B \\ {b} 中参与的基准数量，N_B 是该参考池中的基准总数（即 |B| − 1）。对数项中添加常数 1 以确保数值稳定性。")

# 3.3
add_heading_bilingual(doc, "3.3  Anti-Saturation (SAS)", "3.3  抗饱和度（SAS）", level=2)

add_bilingual_para(doc,
    "Anti-saturation characterizes the structural durability of a benchmark b ∈ B by measuring its remaining performance headroom under the pressure of evolving model capabilities. This dimension is decomposed into two components: Static Weighted Resistance and Dynamic Saturation Projection.",
    "抗饱和度通过衡量基准 b ∈ B 在模型能力演化压力下的剩余性能空间来刻画其结构耐久性。该维度分解为两个组件：静态加权阻力和动态饱和投影。")

# 3.3.1
add_heading_bilingual(doc, "3.3.1  Static Weighted Resistance (SSta)", "3.3.1  静态加权阻力（SSta）", level=3)

add_bilingual_para(doc,
    "SSta estimates a benchmark's difficulty by capability-weighting the score distribution. We utilize θ_i^{(-b)} (explained in Section 3.2) to evaluate benchmark difficulty relative to the strength of its participants. The static resistance is defined as:\nSSta(b) = 1 − Σ_{i∈N} θ_i^{(-b)} · Score_{i,b} / Σ_{i∈N} θ_i^{(-b)}\nwhere Score_{i,b} ∈ [0,1] denotes the normalized performance. This weighted approach applies differential penalties to the remaining headroom, whereby scores from low-capability models induce a more pronounced reduction in SSta while those from frontier models are comparatively preserved. Normalizing by Σ_{i∈N} θ_i^{(-b)} further eliminates the influence of participant volume, ensuring that SSta remains invariant to the total number of evaluated models.",
    "SSta 通过对分数分布进行能力加权来估计基准的难度。我们利用 θ_i^{(-b)}（见第 3.2 节）相对于参与者的强度来评估基准难度。静态阻力定义为：\nSSta(b) = 1 − Σ_{i∈N} θ_i^{(-b)} · Score_{i,b} / Σ_{i∈N} θ_i^{(-b)}\n其中 Score_{i,b} ∈ [0,1] 表示归一化性能。该加权方法对剩余提升空间施加差异化惩罚：低能力模型的分数对 SSta 产生更显著的降低，而前沿模型的分数相对保留。通过 Σ_{i∈N} θ_i^{(-b)} 归一化进一步消除了参与者数量的影响，确保 SSta 对评估模型总数保持不变。")

# 3.3.2
add_heading_bilingual(doc, "3.3.2  Dynamic Saturation Projection (SDyn)", "3.3.2  动态饱和投影（SDyn）", level=3)

add_bilingual_para(doc,
    "SDyn measures the short-term velocity of score inflation by projecting near-future performance under a time-series trend model. We fit a linear trend y = k · t + d using Ordinary Least Squares, where the slope k represents the conquest rate (score increment per day). Using a 30-day projection window, the dynamic saturation is defined as:\nSDyn(b) = max(0, 1 − (Score_b + 30 · k))\nwhere Score_b is the current mean score of benchmark b. If k < 0, typically indicating benchmark difficulty is increasing, SDyn(b) is computed as:\nSAS(b) = 0.8 · SSta + 0.2 · SDyn\nWe assign a lower weight (20%) to SDyn due to the inherent uncertainty of short-term projections, while the static component provides a stable observation of the current saturation status.",
    "SDyn 通过在时间序列趋势模型下投影近期未来性能来衡量分数膨胀的短期速度。我们使用普通最小二乘法拟合线性趋势 y = k · t + d，其中斜率 k 代表征服率（每日分数增量）。使用 30 天投影窗口，动态饱和度定义为：\nSDyn(b) = max(0, 1 − (Score_b + 30 · k))\n其中 Score_b 是基准 b 的当前平均分。如果 k < 0，通常表示基准难度在增加，SDyn(b) 计算为：\nSAS(b) = 0.8 · SSta + 0.2 · SDyn\n由于短期投影的固有不确定性，我们为 SDyn 分配较低的权重（20%），而静态组件提供当前饱和状态的稳定观测。")

# 3.4
add_heading_bilingual(doc, "3.4  Impact (SImp)", "3.4  影响力（SImp）", level=2)

add_bilingual_para(doc,
    "SImp measures the ecological influence of a benchmark b ∈ B through its adoption by capable models. Rather than relying on raw popularity signals such as repository stars, downloads, or historical citation momentum, we define impact as capability-weighted evaluation coverage within the model ecosystem.",
    "SImp 通过高能力模型的采用来衡量基准 b ∈ B 的生态影响力。我们不依赖于原始流行度信号（如仓库星标、下载量或历史引用惯性），而是将影响力定义为模型生态系统中的能力加权评估覆盖度。")

# 3.4.1
add_heading_bilingual(doc, "3.4.1  Capability-weighted Adoption", "3.4.1  能力加权采用度", level=3)

add_bilingual_para(doc,
    "Let M_bench ⊆ M_all denote the set of models that have reported valid evaluation results on benchmark b, where M_all is the full set of models in the dataset. A model is considered to participate in b if it reports a valid score, regardless of the score magnitude. To avoid circularity, we use the LOBO-calibrated capability score defined in Section 3.2, which is computed by excluding benchmark b from the capability estimation process. The raw capability coverage of benchmark b is defined as the square-root mapping of the ratio of capability-weighted participation to total ecosystem capability. This compresses the upper range of the coverage ratio, preventing highly reported benchmarks from dominating solely due to participation volume, while still preserving the advantage of benchmarks evaluated by high-capability models.",
    "设 M_bench ⊆ M_all 表示在基准 b 上报告了有效评估结果的模型集合，其中 M_all 是数据集中的全部模型。如果模型报告了有效分数，则视为参与了 b，无论分数大小。为避免循环性，我们使用第 3.2 节中定义的 LOBO 校准能力分数，该分数在能力估计过程中排除了基准 b。基准 b 的原始能力覆盖度定义为能力加权参与与总生态系统能力之比的平方根映射。这压缩了覆盖比的上限范围，防止高报告量基准仅因参与量而占主导，同时保留了被高能力模型评估的基准的优势。")

# 3.4.2
add_heading_bilingual(doc, "3.4.2  Interpretation", "3.4.2  解读", level=3)

add_bilingual_para(doc,
    "This formulation captures impact as capability-weighted adoption. A benchmark receives a high SImp when a substantial portion of the ecosystem's calibrated model capability is represented among its evaluated participants. Because the capability scores are computed using LOBO calibration, the benchmark being assessed does not contribute to the capability estimates used to measure its own impact. By construction, SImp is bounded within [0,1] and does not depend on benchmark release dates, model release dates, community metadata, or heuristic temporal decay. This design makes the impact score less sensitive to raw popularity accumulation and better aligned with the technical adoption of benchmarks by capable models.",
    "该公式将影响力捕获为能力加权采用度。当生态系统中相当比例的校准模型能力体现在其评估参与者中时，基准获得高 SImp。由于能力分数使用 LOBO 校准计算，被评估的基准不会对用于衡量其自身影响力的能力估计产生贡献。根据构造，SImp 有界于 [0,1]，且不依赖于基准发布日期、模型发布日期、社区元数据或启发式时间衰减。该设计使影响力分数对原始流行度积累不太敏感，更好地与高能力模型对基准的技术采用相一致。")

add_separator(doc)

# ============ 4 DATA PIPELINE ============
add_heading_bilingual(doc, "4  Data Pipeline", "4  数据管线", level=1)

add_bilingual_para(doc,
    "As illustrated in Figure 3, we implement a multi-stage pipeline to distill a high-fidelity benchmark suite from an initial corpus of 416 candidate evaluation sets, systematically aggregated from the technical reports of 117 mainstream models released between 2025 and April 2026. The workflow encompasses rigorous alignment protocols, including a minimum participation threshold for statistical validity, score normalization to a [0,100] range, and a 1.25× headroom calibration for open-ended metrics to mitigate potential ceiling effects. After specialized handling for temporal snapshots and tool-augmented benchmarks, we retain a final set of 158 validated benchmarks spanning diverse domains and vendors. For brevity, the exhaustive details regarding the filtering, normalization, and calibration methodologies are provided in Appendix E. The complete metadata is publicly available at our anonymous repository.",
    "如图 3 所示，我们实施了一个多阶段管线，从 416 个候选评估集的初始语料中提炼出高保真基准套件，这些候选集系统性地从 2025 年至 2026 年 4 月期间发布的 117 个主流模型技术报告中聚合而成。工作流程包含严格的校准协议，包括用于统计有效性的最低参与阈值、分数归一化至 [0,100] 范围，以及针对开放指标的 1.25× 提升空间校准以缓解潜在的天花板效应。经过对时间快照和工具增强基准的专门处理，我们保留了 158 个跨多元领域和供应商的验证基准。关于筛选、归一化和校准方法的详尽细节见附录 E。完整元数据可在我们的匿名仓库中公开获取。")

add_separator(doc)

# ============ 5 RESULTS AND VALIDATION ============
add_heading_bilingual(doc, "5  Results and Validation", "5  结果与验证", level=1)

# 5.1
add_heading_bilingual(doc, "5.1  Main BHI Results", "5.1  主要 BHI 结果", level=2)

add_bilingual_para(doc,
    "The final BHI is synthesized from the three primary dimensions introduced in Section 3: SDisc, SAS, and SImp. In this process, the SAS and SImp axes are adjusted by the Model Capability Calibration component—derived from the LOBO-adjusted win rate—to transform heterogeneous benchmark-level signals into a unified health score. We utilize the CRITIC objective weighting framework for the final index aggregation to maintain a consistent data-driven formulation. This method determines the importance of each dimension by evaluating its internal information volume, measured by contrast intensity, and its degree of independence, measured by conflicting character, within the global distribution of 158 benchmarks. Under this framework, the optimized weighting distribution is 32.63% for SDisc, 35.74% for SAS, and 31.63% for SImp, resulting in the final formulation:\nBHI(b) = w_Disc · SDisc(b) + w_AS · SAS(b) + w_Imp · SImp(b)",
    "最终 BHI 由第 3 节介绍的三个主要维度合成：SDisc、SAS 和 SImp。在此过程中，SAS 和 SImp 轴由模型能力校准组件调整——该组件源自 LOBO 调整胜率——将异构的基准级信号转换为统一的健康分数。我们利用 CRITIC 客观加权框架进行最终指数聚合，以保持一致的数据驱动公式。该方法通过评估每个维度的内部信息量（由对比强度衡量）和独立程度（由冲突特征衡量）来确定其重要性。在该框架下，优化权重分配为 SDisc 32.63%、SAS 35.74% 和 SImp 31.63%，最终公式为：\nBHI(b) = w_Disc · SDisc(b) + w_AS · SAS(b) + w_Imp · SImp(b)")

add_bilingual_para(doc,
    "Representative scores and rankings are summarized in Table 1, while the comprehensive rankings for the complete set of 158 benchmarks are provided in Table 12. Overall, the resulting ranking reflects a balance between technical resolution, saturation resistance, and capability-weighted adoption, rather than optimizing for any single benchmark property.",
    "代表性分数和排名汇总于表 1，158 个基准的完整排名见表 12。总体而言，结果排名反映了技术分辨率、抗饱和度和能力加权采用之间的平衡，而非优化任何单一基准属性。")

# 5.2
add_heading_bilingual(doc, "5.2  Component Necessity: Ablation Study", "5.2  组件必要性：消融研究", level=2)

add_bilingual_para(doc,
    "To assess the necessity of each component in the BHI formulation, we conduct an ablation study by removing one indicator at a time and recomputing the final ranking. The ranking differences are measured relative to the full model using average and maximum absolute rank shifts. Removing SAS yields the largest average rank shift, highlighting its importance in identifying unsaturated benchmarks. Removing SDisc produces a comparable average shift, highlighting the importance of discriminative resolution. Although removing SImp yields a smaller average shift, it still introduces significant benchmark-level corrections, confirming that capability-weighted adoption contributes non-trivially to the final ranking. These results demonstrate that all three components are necessary and jointly define the BHI framework.",
    "为评估 BHI 公式中每个组件的必要性，我们进行了消融研究，每次移除一个指标并重新计算最终排名。排名差异相对于完整模型使用平均和最大绝对排名偏移来衡量。移除 SAS 产生最大的平均排名偏移，突出了其在识别未饱和基准中的重要性。移除 SDisc 产生相当的平均偏移，突出了区分分辨率的重要性。虽然移除 SImp 产生较小的平均偏移，但它仍引入显著的基准级修正，确认能力加权采用对最终排名有非平凡贡献。这些结果表明三个组件都是必要的，共同定义了 BHI 框架。")

# 5.3
add_heading_bilingual(doc, "5.3  Component Complementarity: Indicator Orthogonality", "5.3  组件互补性：指标正交性", level=2)

add_bilingual_para(doc,
    "We further analyze the relationships between the three indicators using both Pearson and Spearman correlation coefficients. The correlation analysis shows that the indicators are complementary rather than redundant. The strongest association occurs between SDisc and SAS (Spearman ρ = 0.46), which is expected since highly saturated benchmarks tend to exhibit reduced discriminative resolution. However, all pairwise correlations remain below 0.5, and SImp shows consistently weak correlations with both technical indicators. This indicates that SImp captures an adoption-oriented signal distinct from benchmark difficulty and discrimination, while SDisc and SAS provide related but non-identical views of benchmark quality. Overall, the results confirm that the three components encode complementary aspects of benchmark health.",
    "我们进一步使用 Pearson 和 Spearman 相关系数分析三个指标之间的关系。相关性分析表明指标是互补的而非冗余的。最强的关联出现在 SDisc 和 SAS 之间（Spearman ρ = 0.46），这是预期的，因为高度饱和的基准倾向于表现出降低的区分分辨率。然而，所有成对相关性保持在 0.5 以下，SImp 与两个技术指标均显示一致弱相关。这表明 SImp 捕获了与基准难度和区分度不同的以采用为导向的信号，而 SDisc 和 SAS 提供了基准质量的关联但非同一视角。总体而言，结果确认三个组件编码了基准健康的互补方面。")

add_bilingual_para(doc,
    "Robustness and Sensitivity. We evaluate the robustness of the proposed framework under multiple perturbations, including score noise, model dropout, parameter variation, and aggregation schemes. Under Gaussian noise (σ ∈ [0.01,0.2]), the resulting rankings remain stable, with Spearman correlations remaining above 0.9 under moderate perturbation and degrading only gradually as the noise level increases. Similarly, random model dropout from 5% to 60% produces only gradual ranking degradation, indicating resilience to incomplete evaluation coverage.",
    "稳健性与敏感性。我们在多种扰动下评估所提出框架的稳健性，包括分数噪声、模型缺失、参数变化和聚合方案。在高斯噪声（σ ∈ [0.01,0.2]）下，结果排名保持稳定，中等扰动下 Spearman 相关性保持在 0.9 以上，仅在噪声水平增加时逐渐退化。类似地，5% 到 60% 的随机模型缺失仅产生渐进的排名退化，表明对不完整评估覆盖具有韧性。")

add_bilingual_para(doc,
    "We further analyze sensitivity to key design choices in the anti-saturation component. Varying the static/dynamic weighting in SAS and the temporal projection window used in SDyn produces negligible changes in the final ranking, with Spearman correlations above 0.99 across all tested configurations. We also compare the CRITIC objective weighting scheme with uniform weighting and observe that CRITIC yields consistently more stable rankings under both noise injection and model dropout, supporting its use as a data-driven aggregation strategy.",
    "我们进一步分析了抗饱和度组件中关键设计选择的敏感性。改变 SAS 中的静态/动态加权和 SDyn 中使用的时间投影窗口对最终排名产生可忽略的变化，所有测试配置下 Spearman 相关性均高于 0.99。我们还比较了 CRITIC 客观加权方案与均匀加权，观察到 CRITIC 在噪声注入和模型缺失下均产生更稳定的排名，支持其作为数据驱动聚合策略的使用。")

add_bilingual_para(doc,
    "To examine potential bias from data availability, we analyze the relationship between benchmark coverage and BHI scores. Although coverage is strongly correlated with SImp by construction, its correlation with SDisc and SAS remains weak, and the overall correlation with BHI is moderate, indicating that the final ranking is not dominated by benchmark popularity or evaluation frequency. Overall, these results demonstrate that the BHI framework is robust to score noise, model sparsity, parameter choices, and coverage variation.",
    "为检验数据可用性的潜在偏差，我们分析了基准覆盖度与 BHI 分数之间的关系。虽然覆盖度与 SImp 在构造上强相关，但其与 SDisc 和 SAS 的相关性保持弱，与 BHI 的整体相关性为中等，表明最终排名不被基准流行度或评估频率所主导。总体而言，这些结果表明 BHI 框架对分数噪声、模型稀疏性、参数选择和覆盖度变化具有稳健性。")

add_separator(doc)

# ============ 6 EXTERNAL VALIDATION ============
add_heading_bilingual(doc, "6  External Validation and Case Study", "6  外部验证与案例研究", level=1)

add_bilingual_para(doc,
    "To examine whether BHI is consistent with external evidence, we conduct a diagnostic case study using two benchmark cohorts defined independently of the BHI computation. The known-saturated cohort includes benchmarks widely regarded as losing frontier-level diagnostic value due to ceiling effects, score compression, leakage concerns, or replacement by stronger successors, while the active frontier cohort includes benchmarks that remain in recent frontier-model reports for reasoning, browsing, coding, and agentic evaluation. These cohort labels are used only for validation and are never used to compute BHI scores.",
    "为检验 BHI 是否与外部证据一致，我们使用两个独立于 BHI 计算定义的基准队列进行了诊断性案例研究。已知饱和队列包括因天花板效应、分数压缩、泄漏问题或被更强继任者替代而广泛认为失去前沿诊断价值的基准，而活跃前沿队列包括在近期前沿模型报告（推理、浏览、编码和智能体评估）中仍出现的基准。这些队列标签仅用于验证，从不用于计算 BHI 分数。")

add_bilingual_para(doc,
    "Figure 6 provides a temporal check: known-saturated benchmarks show reduced presence in recent frontier reports, whereas active frontier benchmarks remain in contemporary evaluation practice. Table 3 further shows a clear separation between the two groups, with known-saturated benchmarks averaging 0.227 ± 0.048 and active frontier benchmarks averaging 0.600 ± 0.038. This supports directional consistency with external practice, though it should be interpreted as face validity rather than calibration against an absolute ground-truth health score.",
    "图 6 提供了时间检查：已知饱和基准在近期前沿报告中的出现减少，而活跃前沿基准仍在当代评估实践中出现。表 3 进一步显示两组之间的清晰分离，已知饱和基准平均 0.227 ± 0.048，活跃前沿基准平均 0.600 ± 0.038。这支持了与外部实践的方向一致性，尽管这应被理解为表面效度而非针对绝对真实健康分数的校准。")

add_bilingual_para(doc,
    "The component decomposition further explains why benchmarks receive low BHI scores. HellaSwag and ARC-Challenge are mainly penalized by severe ceiling effects, reflected in extremely low SAS. HumanEval and MBPP retain some discriminative signal but have low anti-saturation scores, suggesting that function-level coding tasks have largely been absorbed by contemporary training and evaluation practice. C-Eval degrades across all three axes, consistent with its reduced reliability as a modern subject-knowledge benchmark. Importantly, GSM8K and MMLU still retain non-trivial SImp, indicating that their low BHI scores are not merely caused by low visibility or sparse reporting, but are mainly saturation-driven. This decomposition shows that BHI not only ranks benchmarks, but also distinguishes failure modes including saturation-driven obsolescence, weak adoption, and reduced discrimination.",
    "组件分解进一步解释了基准为何获得低 BHI 分数。HellaSwag 和 ARC-Challenge 主要因严重的天花板效应而受罚，反映在极低的 SAS 中。HumanEval 和 MBPP 保留了一些区分信号，但抗饱和度分数低，表明函数级编码任务在很大程度上已被当代训练和评估实践吸收。C-Eval 在所有三个轴上退化，与其作为现代学科知识基准的可靠性降低一致。重要的是，GSM8K 和 MMLU 仍保留非平凡的 SImp，表明其低 BHI 分数并非仅仅由低可见度或稀疏报告造成，而主要是饱和驱动的。该分解表明 BHI 不仅对基准排名，还能区分失效模式，包括饱和驱动的过时、弱采用和区分度降低。")

add_bilingual_para(doc,
    "Overall, these case studies support BHI as a data-driven benchmark auditing framework whose outputs align with external evaluation practice while remaining interpretable at the component level. More detailed domain-specific analyses are provided in Appendix J.",
    "总体而言，这些案例研究支持 BHI 作为数据驱动的基准审计框架，其输出与外部评估实践一致，同时在组件层面保持可解释性。更详细的领域特定分析见附录 J。")

add_separator(doc)

# ============ 7 CONCLUSION ============
add_heading_bilingual(doc, "7  Conclusion", "7  结论", level=1)

add_bilingual_para(doc,
    "This paper introduces the Benchmark Health Index (BHI), a data-driven framework for auditing LLM benchmarks at the ecosystem level. Unlike conventional model-centric leaderboards, BHI treats benchmarks themselves as objects of evaluation and quantifies their health through discrimination, anti-saturation, and capability-weighted adoption. Based on large-scale public technical reports, BHI produces interpretable benchmark rankings and reveals current failure modes in the evaluation ecosystem. Validation further shows that its components are necessary, complementary, and stable. Overall, BHI provides a reproducible quantitative basis for benchmark selection, lifecycle monitoring, and systematic benchmark governance.",
    "本文提出了基准健康指数（BHI），一个在生态系统层面审计 LLM 基准的数据驱动框架。与传统的以模型为中心的排行榜不同，BHI 将基准本身视为评估对象，并通过区分度、抗饱和度和能力加权采用来量化其健康状况。基于大规模公开技术报告，BHI 产生了可解释的基准排名，并揭示了评估生态系统中当前的失效模式。验证进一步表明其组件是必要的、互补的和稳定的。总体而言，BHI 为基准选择、生命周期监测和系统性基准治理提供了可复现的量化基础。")

add_separator(doc)

# ============ APPENDIX A ============
add_heading_bilingual(doc, "A  Limitations", "A  局限性", level=1)

add_bilingual_para(doc,
    "BHI is, to the best of our knowledge, the first benchmark-level auditing framework that systematically quantifies benchmark health across discrimination, anti-saturation, and ecosystem impact, rather than an absolute measure of model capability or benchmark quality. Its conclusions should therefore be interpreted as relative diagnostic signals within the collected evaluation ecosystem. The framework is most informative when a benchmark has sufficient public reporting coverage and when the reported scores are comparable after normalization. Benchmarks with limited adoption, highly specialized evaluation protocols, or rapidly changing versions may receive less stable estimates, even though the LOBO calibration in Section 3.2 and the coverage analysis in Appendix I are designed to reduce this effect.",
    "据我们所知，BHI 是首个系统性量化基准在区分度、抗饱和度和生态影响力方面的健康状况的基准级审计框架，而非模型能力或基准质量的绝对度量。因此，其结论应被解释为所收集评估生态系统内的相对诊断信号。当基准具有足够的公开报告覆盖度且报告的分数在归一化后可比较时，该框架最具信息量。采用有限、评估协议高度专业化或版本快速变更的基准可能获得不太稳定的估计，尽管第 3.2 节的 LOBO 校准和附录 I 的覆盖度分析旨在减少这种影响。")

add_bilingual_para(doc,
    "A second limitation comes from the heterogeneity of public technical reports. Although the data pipeline applies systematic filtering, score normalization, and benchmark alignment, model reports do not always disclose evaluation prompts, decoding settings, test splits, contamination controls, or metric variants with the same level of detail. As a result, BHI should not be viewed as replacing benchmark-specific replication studies. Instead, it provides a reproducible macro-level audit that can help identify where such more detailed studies are most needed.",
    "第二个局限性来自公开技术报告的异构性。虽然数据管线应用了系统性筛选、分数归一化和基准校准，但模型报告并不总是以相同的详细程度披露评估提示、解码设置、测试分割、污染控制或指标变体。因此，BHI 不应被视为替代基准特定的复现研究。相反，它提供了可复现的宏观层面审计，可帮助识别最需要进行更详细研究的领域。")

add_bilingual_para(doc,
    "The anti-saturation component also relies on several simplifying assumptions. We model short-term score evolution using a linear trend and a fixed projection window. This choice is intentionally conservative and is supported by the sensitivity analyses, but real benchmark lifecycles may be nonlinear: a benchmark can experience abrupt saturation after a new training recipe, tool-use capability, or data exposure event becomes widespread. Therefore, the dynamic component should be interpreted as a near-term warning signal rather than a long-horizon forecast.",
    "抗饱和度组件也依赖于若干简化假设。我们使用线性趋势和固定投影窗口建模短期分数演化。这一选择是有意保守的，并得到敏感性分析的支持，但实际基准生命周期可能是非线性的：基准可能在新的训练方法、工具使用能力或数据暴露事件普及后经历突然饱和。因此，动态组件应被解释为近期预警信号而非长期预测。")

add_bilingual_para(doc,
    "The impact score measures capability-weighted adoption among models with valid public reports. This design avoids relying on raw popularity signals such as citations or repository statistics, but it also means that BHI captures reported evaluation practice rather than all forms of scientific influence. For example, a benchmark may shape research discussions, diagnostic analysis, or internal evaluation pipelines without appearing frequently in public model reports. Conversely, a widely reported benchmark may retain high adoption even after its technical diagnostic value has begun to decline. The three-axis decomposition is intended to expose this tension, but it cannot fully observe private or undocumented evaluation use.",
    "影响力分数衡量具有有效公开报告的模型中的能力加权采用度。该设计避免依赖引用或仓库统计等原始流行度信号，但这也意味着 BHI 捕获的是报告的评估实践而非所有形式的科学影响力。例如，基准可能塑造研究讨论、诊断分析或内部评估管线，而无需频繁出现在公开模型报告中。反之，广泛报告的基准可能在其技术诊断价值开始下降后仍保持高采用度。三轴分解旨在揭示这种张力，但无法完全观察私有或未记录的评估使用。")

add_bilingual_para(doc,
    "Our external validation is also necessarily limited. The case study shows that BHI separates independently identified saturated benchmarks from active frontier benchmarks, providing evidence of directional consistency with current evaluation practice. However, this should be understood as face validity rather than criterion validity against an independently collected ground-truth health score. Building such a ground truth would require structured expert elicitation or longitudinal replication studies, which we leave for future work.",
    "我们的外部验证也必然有限。案例研究表明 BHI 能将独立识别的饱和基准与活跃前沿基准分离，提供了与当前评估实践方向一致性的证据。然而，这应被理解为表面效度而非针对独立收集的真实健康分数的准则效度。建立这样的真实标准需要结构化专家征询或纵向复现研究，我们将其留给未来工作。")

add_bilingual_para(doc,
    "Finally, BHI is a lightweight post-hoc auditing procedure, but its scalability still depends on the availability and maintenance of structured benchmark metadata. The computational cost of the index itself is modest compared with model evaluation, since it operates on reported scores rather than rerunning benchmarks. The main cost lies in data curation, version alignment, and continuous updates as new models and benchmarks appear. Future work could integrate BHI with existing agent frameworks to support automated benchmark report parsing, metadata alignment, and periodic index updates as new models and benchmarks emerge.",
    "最后，BHI 是一个轻量级的事后审计程序，但其可扩展性仍依赖于结构化基准元数据的可用性和维护。指数本身的计算成本与模型评估相比较低，因为它操作的是报告的分数而非重新运行基准。主要成本在于数据策划、版本对齐以及随新模型和基准出现而进行的持续更新。未来工作可以将 BHI 与现有的智能体框架集成，以支持自动化基准报告解析、元数据对齐以及随新模型和基准出现的定期指数更新。")

add_separator(doc)

# ============ APPENDIX B ============
add_heading_bilingual(doc, "B  Expanded Related Work", "B  扩展相关工作", level=1)

add_heading_bilingual(doc, "B.1  Shifting Evaluation Paradigms", "B.1  评估范式的转变", level=2)

add_bilingual_para(doc,
    "The meteoric rise of LLMs has catalyzed a paradigm shift in evaluation, moving from single-task leaderboards toward comprehensive benchmark suites that encompass knowledge, reasoning, and safety. However, conventional evaluation frameworks predominantly adopt a model-centric perspective, treating benchmarks as static and immutable measurement tools. As model capabilities continue to leapfrog, public benchmarks are increasingly susceptible to score inflation and leaderboard distortion, which diminish their capacity to characterize genuine differences in model performance. Consequently, the research focus in the community is gradually pivoting from simple performance comparisons between models to in-depth investigations into the validity, robustness, and lifecycles of the benchmarks themselves. Despite this shift, existing studies typically propose diagnostic metrics centered on local dimensions—such as uncertainty, measurement quality, or cross-benchmark consistency—and lack a unified, scalable framework for the longitudinal monitoring of benchmark efficacy as it co-evolves with model advancements.",
    "LLM 的迅猛崛起催化了评估范式的转变，从单任务排行榜转向涵盖知识、推理和安全性的综合基准套件。然而，传统评估框架主要采用以模型为中心的视角，将基准视为静态且不可变的测量工具。随着模型能力持续跨越，公开基准越来越容易受到分数膨胀和排行榜扭曲的影响，这削弱了其刻画模型性能真实差异的能力。因此，社区的研究重点正逐渐从模型间的简单性能比较转向对基准本身有效性、稳健性和生命周期的深入调查。尽管有这种转变，现有研究通常提出集中于局部维度的诊断指标——如不确定性、测量质量或跨基准一致性——缺乏一个统一的、可扩展的框架来纵向监测基准效力随模型进步而协同演化的情况。")

add_heading_bilingual(doc, "B.2  Reliability and Measurement Theory", "B.2  可靠性与测量理论", level=2)

add_bilingual_para(doc,
    "The reliability of evaluation outcomes faces severe challenges from prompt sensitivity, annotation noise, and stochasticity in sampling. To mitigate these issues, researchers have integrated statistical tools such as bootstrap estimation, confidence intervals, and significance testing to distinguish true performance gains from measurement noise, establishing more robust reporting norms. From the perspectives of psychometrics and Item Response Theory (IRT), the core utility of a benchmark lies in its discriminative power—the measurement resolution jointly determined by item difficulty and discrimination parameters. In the context of LLM evaluation, however, high-performing models often trigger upper-tail compression and a lack of leaderboard separability, rendering traditional metrics ineffective at identifying fine-grained capability nuances.",
    "评估结果的可靠性面临来自提示敏感性、标注噪声和采样随机性的严峻挑战。为缓解这些问题，研究人员整合了自助估计、置信区间和显著性检验等统计工具，以区分真实的性能增益与测量噪声，建立了更稳健的报告规范。从心理测量学和项目反应理论（IRT）的角度，基准的核心效用在于其区分力——由项目难度和区分参数共同决定的测量分辨率。然而，在 LLM 评估的背景下，高性能模型常常引发上尾压缩和排行榜可分性不足，使传统指标在识别细粒度能力差异方面变得无效。")

add_heading_bilingual(doc, "B.3  Benchmark Degradation and Dynamics", "B.3  基准退化与动态", level=2)

add_bilingual_para(doc,
    "The efficacy of evaluation benchmarks exhibits a discernible degradation trend over time, primarily driven by saturation effects and data contamination. When models approach ceiling performance on existing datasets, or when evaluation data is implicitly leaked into pre-training corpora, the discriminative capacity of a benchmark shrinks rapidly. This blurs the boundary between generalization and memorization, prompting the development of contamination detection and verifiable evaluation corrections. To address the loss of resolution caused by saturation, some work attempts to reactivate the discriminative power of existing benchmarks through weighted or augmented metrics, thus improving the separability among SOTA models. Although dynamic evaluation—which refreshes the task pools or introduces temporal windows—can alleviate contamination risks, it introduces a trade-off between maintaining longitudinal comparability and controlling evaluation costs. Therefore, a significant research gap remains in the construction of a unified data-driven framework for the long-term tracking of benchmark health.",
    "评估基准的效力随时间呈现出可辨识的退化趋势，主要由饱和效应和数据污染驱动。当模型在现有数据集上接近天花板性能，或评估数据隐含泄漏到预训练语料中时，基准的区分能力迅速缩小。这模糊了泛化与记忆之间的界限，推动了污染检测和可验证评估修正的发展。为解决饱和导致的分辨率损失，一些工作尝试通过加权或增强指标重新激活现有基准的区分力，从而改善 SOTA 模型之间的可分性。虽然动态评估——刷新任务池或引入时间窗口——可以缓解污染风险，但它引入了维持纵向可比性和控制评估成本之间的权衡。因此，在构建用于长期追踪基准健康的统一数据驱动框架方面，仍然存在显著的研究空白。")

add_separator(doc)

# ============ APPENDIX C ============
add_heading_bilingual(doc, "C  Sensitivity Analysis of Discrimination Threshold", "C  区分阈值的敏感性分析", level=1)

add_bilingual_para(doc,
    "The Effective Differentiation Ratio (EDR) in Section 3.1.1 depends on a noise-filtering threshold δ, which determines whether a pairwise score gap is treated as an informative distinction between two models. In the main experiments, we set δ = 2.0% of the observed score range for each benchmark. This appendix evaluates whether the resulting BHI ranking is sensitive to this choice. We vary δ from 0.5% to 5.0% and recompute the complete BHI pipeline under each setting. The resulting ranking is compared against the default setting δ = 2.0% using Spearman rank correlation and top-10 overlap.",
    "第 3.1.1 节中的有效区分比（EDR）依赖于噪声过滤阈值 δ，该阈值决定成对分数差距是否被视为两个模型之间的信息性区分。在主要实验中，我们设置 δ = 2.0% 的观察分数范围。本附录评估 BHI 排名是否对该选择敏感。我们将 δ 从 0.5% 变化到 5.0%，并在每个设置下重新计算完整的 BHI 管线。结果排名使用 Spearman 排名相关性和前 10 重叠与默认设置 δ = 2.0% 进行比较。")

add_bilingual_para(doc,
    "As shown in Table 4, the BHI ranking remains highly stable across the tested threshold range. Even under the most conservative threshold setting of δ = 5.0%, the Spearman correlation with the default ranking remains 0.9646, and the top-10 overlap remains 9 out of 10. The strongest stability is observed near the default region between 1.5% and 2.5%, where the Spearman correlation exceeds 0.995. Figure 7 further shows that the CRITIC-assigned weight of SDisc varies smoothly with δ. Larger thresholds make EDR stricter and increase the relative information contribution of the discrimination axis, but this does not materially alter the final ranking. These results indicate that the proposed discrimination component captures a stable ranking signal rather than an artifact of a specific threshold choice.",
    "如表 4 所示，BHI 排名在测试阈值范围内保持高度稳定。即使在最保守的阈值设置 δ = 5.0% 下，与默认排名的 Spearman 相关性仍为 0.9646，前 10 重叠为 9/10。最强的稳定性观察到在默认区域 1.5% 至 2.5% 之间，Spearman 相关性超过 0.995。图 7 进一步显示 CRITIC 分配的 SDisc 权重随 δ 平滑变化。较大的阈值使 EDR 更严格，增加区分轴的相对信息贡献，但这不会实质性地改变最终排名。这些结果表明所提出的区分组件捕获的是稳定的排名信号，而非特定阈值选择的产物。")

add_separator(doc)

# ============ APPENDIX D ============
add_heading_bilingual(doc, "D  SDM Weighting and Normalization", "D  SDM 加权与归一化", level=1)

add_bilingual_para(doc,
    "This subsection provides the detailed mathematical derivation for SDM and the normalization process utilized in computing SDisc(b).",
    "本小节提供了 SDM 和计算 SDisc(b) 时使用的归一化过程的详细数学推导。")

add_heading_bilingual(doc, "D.1  Analysis of Metric Characteristics and Scale Disparities", "D.1  指标特征与尺度差异分析", level=2)

add_bilingual_para(doc,
    "When constructing SDisc(b), the two constituent metrics, EDR(b) and RCV(b), exhibit significant differences in their mathematical properties and distributional behaviors:\n1. Incongruent Scales and Units: RCV(b) is derived from the percentile spread of model score distributions and typically manifests as a continuous numerical value. In contrast, EDR(b) represents the proportion of model pairs exceeding a specific noise-filtering threshold.\n2. Variance Suppression Induced by Discretization: The values of EDR(b) are highly sensitive to the number of participating models N. For instance, if a benchmark has only N = 3 models, EDR(b) can only take values from the discrete set {0,1/3,2/3,1}. Such high discreteness in sparse datasets can lead to a pseudo-standard deviation that is statistically lower than the metric's true discriminative potential.\n3. Sparse Coverage Bias: Given that the study analyzes a meta-dataset of 158 benchmarks with varying model coverage, using raw standard deviations for weighting would allow metrics with naturally wider or more continuous distributions, such as RCV(b), to dominate the final weight, thereby masking the fine-grained micro-resolution signals provided by EDR(b).",
    "在构建 SDisc(b) 时，两个组成指标 EDR(b) 和 RCV(b) 在数学性质和分布行为上表现出显著差异：\n1. 尺度和单位不一致：RCV(b) 源自模型分数分布的百分位扩展，通常表现为连续数值。相比之下，EDR(b) 代表超过特定噪声过滤阈值的模型对比例。\n2. 离散化引起的方差抑制：EDR(b) 的值对参与模型数量 N 高度敏感。例如，如果基准只有 N = 3 个模型，EDR(b) 只能取离散集合 {0,1/3,2/3,1} 中的值。这种在稀疏数据集中的高离散性可能导致伪标准差在统计上低于指标的真实区分潜力。\n3. 稀疏覆盖偏差：鉴于研究分析了 158 个模型覆盖度不同的基准的元数据集，使用原始标准差进行加权将允许自然分布更宽或更连续的指标（如 RCV(b)）主导最终权重，从而掩盖 EDR(b) 提供的细粒度微观分辨率信号。")

add_heading_bilingual(doc, "D.2  Scale Alignment via Min-Max Normalization", "D.2  通过最小-最大归一化进行尺度对齐", level=2)

add_bilingual_para(doc,
    "To mitigate the aforementioned biases, we first perform Min-Max normalization. For any indicator k ∈ {EDR,RCV}, the normalized value Norm(k(b)) is calculated as:\nNorm(k(b)) = (k(b) − min(k)) / (max(k) − min(k))\nwhere min(k) and max(k) denote the minimum and maximum values of the metric across the entire set of 158 validated benchmarks. By mapping these indicators to the [0,1] interval, we effectively decouple the distributional shape from the magnitude scale. Within this normalized space, any observed fluctuation, represented by the standard deviation, solely reflects the contrast intensity of the metric in evaluating benchmark health, rather than being an artifact of the raw calculation formula or sampling density.",
    "为缓解上述偏差，我们首先执行最小-最大归一化。对于任何指标 k ∈ {EDR,RCV}，归一化值 Norm(k(b)) 计算为：\nNorm(k(b)) = (k(b) − min(k)) / (max(k) − min(k))\n其中 min(k) 和 max(k) 表示指标在所有 158 个验证基准中的最小值和最大值。通过将这些指标映射到 [0,1] 区间，我们有效地将分布形状与量级尺度解耦。在这个归一化空间中，任何观察到的波动（由标准差表示）仅反映指标在评估基准健康方面的对比强度，而非原始计算公式或采样密度的产物。")

add_heading_bilingual(doc, "D.3  Mathematical Derivation of SDM Weighting", "D.3  SDM 加权的数学推导", level=2)

add_bilingual_para(doc,
    "Within the normalized coordinate system, we employ SDM to determine the weights w_k, aiming to amplify signals with higher information load:\n1. Calculation of Normalized Standard Deviation: For each indicator k ∈ {EDR,RCV}, we compute its standard deviation σ_k across all benchmarks.\n2. Weight Allocation: The weight w_k is determined by the proportion of the indicator's standard deviation relative to the total: w_k = σ_k / Σ_{k'} σ_{k'}\nThis mechanism ensures that if a metric remains more dispersed after normalization, it demonstrates higher sensitivity to the varying health statuses of different benchmarks and is thus assigned a higher weight.",
    "在归一化坐标系中，我们使用 SDM 确定权重 w_k，旨在放大具有更高信息量的信号：\n1. 归一化标准差的计算：对于每个指标 k ∈ {EDR,RCV}，我们计算其在所有基准上的标准差 σ_k。\n2. 权重分配：权重 w_k 由指标标准差相对于总量的比例确定：w_k = σ_k / Σ_{k'} σ_{k'}\n该机制确保如果指标在归一化后仍然更分散，则表明它对不同基准的不同健康状态具有更高的敏感性，因此被分配更高的权重。")

add_separator(doc)

# ============ APPENDIX E ============
add_heading_bilingual(doc, "E  Expanded Data Pipeline", "E  扩展数据管线", level=1)

add_bilingual_para(doc,
    "BHI employs a multi-stage pipeline to filter, standardize, and calibrate results from heterogeneous evaluation sets. As illustrated in Figure 3, the workflow refines a broad corpus into a high-fidelity benchmark through three stages: collection, alignment, and validation.",
    "BHI 采用多阶段管线来筛选、标准化和校准来自异构评估集的结果。如图 3 所示，工作流程通过三个阶段将广泛的语料精炼为高保真基准：收集、校准和验证。")

add_heading_bilingual(doc, "E.1  Data Collection", "E.1  数据收集", level=2)

add_bilingual_para(doc,
    "We screened the official technical reports of 117 representative models released between 2025 and April 2026 to aggregate an initial corpus of 416 candidate benchmarks. Among them, 91 models were released in 2025, whose chronological distribution is shown in Figure 8. This selection covers both the 2025 model ecosystem and early-2026 frontier releases, providing sufficient coverage of generalized and specialized model capabilities.",
    "我们筛选了 2025 年至 2026 年 4 月期间发布的 117 个代表性模型的官方技术报告，聚合了 416 个候选基准的初始语料。其中，91 个模型于 2025 年发布，其时间分布如图 8 所示。该选择覆盖了 2025 年模型生态系统和 2026 年初的前沿发布，提供了对通用和专业模型能力的充分覆盖。")

add_heading_bilingual(doc, "E.2  Data Alignment", "E.2  数据校准", level=2)

add_bilingual_para(doc,
    "To ensure comparability across disparate metrics and varying evaluation conditions, we implement a multi-stage alignment protocol. This process standardizes raw scores and addresses the structured requirements of different benchmarks.",
    "为确保跨不同指标和变化评估条件的可比性，我们实施了多阶段校准协议。该过程标准化原始分数并处理不同基准的结构化需求。")

add_heading_bilingual(doc, "E.2.1  Filtering, Normalization and Calibration", "E.2.1  筛选、归一化与校准", level=3)

add_bilingual_para(doc,
    "We apply a Minimum Participation Threshold to exclude benchmarks evaluated by fewer than three models, ensuring the statistical validity of our dispersion-based metrics. Subsequently, all scores are normalized to a [0,100] scale. For cost-oriented metrics where lower values indicate better performance, we apply a subtraction-based transformation to align them with a standard reward-centric orientation. For open-ended metrics lacking a theoretical maximum—such as task-specific throughput or reward-modeling scores—we apply a Headroom Buffer Factor of 1.25. The normalized score S_norm is calculated as: S_norm = S_raw / (1.25 · S_max,obs) × 100, where S_max,obs represents the maximum observed performance within the dataset. This maps the current SOTA performance to 80 points, reserving a 20% headroom buffer. This mechanism effectively mitigates ceiling effects and stabilizes saturation diagnostics as model performance continues to evolve.",
    "我们应用最低参与阈值排除少于三个模型评估的基准，确保基于分散度指标的统计有效性。随后，所有分数归一化至 [0,100] 尺度。对于值越低表示性能越好的成本导向指标，我们应用减法变换将其与标准奖励导向对齐。对于缺乏理论最大值的开放指标——如任务特定吞吐量或奖励建模分数——我们应用 1.25 的提升空间缓冲因子。归一化分数 S_norm 计算为：S_norm = S_raw / (1.25 · S_max,obs) × 100，其中 S_max,obs 表示数据集中的最大观察性能。这将当前 SOTA 性能映射到 80 分，保留 20% 的提升空间缓冲。该机制有效缓解了天花板效应，并在模型性能持续演化时稳定饱和诊断。")

add_heading_bilingual(doc, "E.2.2  Specialized Handling", "E.2.2  专门处理", level=3)

add_bilingual_para(doc,
    "For temporal datasets such as LiveCodeBench, we aggregated fragmented, time-bound snapshots into major version clusters under the assumption of difficulty invariance. This consolidation maintains a dense evaluation matrix while preserving the temporal integrity of the data. For tool-augmented evaluations, we implemented a bifurcated selection policy: we utilized base-ability scores for Capability Discrimination to isolate the benchmark's intrinsic resolution in separating model architectures, while incorporating all performance gains—including agentic and retrieval-based optimizations—for Anti-Saturation to evaluate the practical exhaustion of the task environment.",
    "对于时间数据集（如 LiveCodeBench），我们在难度不变性假设下将碎片化的、有时间限制的快照聚合为主要版本集群。这种整合在保留数据时间完整性的同时维持了密集的评估矩阵。对于工具增强的评估，我们实施了分叉选择策略：我们利用基础能力分数进行能力区分，以隔离基准在分离模型架构方面的内在分辨率；同时纳入所有性能增益——包括智能体和基于检索的优化——用于抗饱和度，以评估任务环境的实际耗竭程度。")

add_heading_bilingual(doc, "E.3  Validated Benchmark Set", "E.3  验证基准集", level=2)

add_bilingual_para(doc,
    "After rigorous alignment and filtering, we retained 158 validated benchmarks. The domain distribution highlights BHI's diverse evaluation dimensions, covering core domains like Reasoning (20.0%) and Multimodal (14.8%) alongside Mathematics and Code generation. Moreover, the vendor distribution reflects ecosystem diversity, spanning closed-source (OpenAI, Anthropic), open-source (DeepSeek, Alibaba), and domain-specific models (Meituan, Xiaomi). This curated suite establishes a high-fidelity foundation for statistical analysis and BHI health metric computation.",
    "经过严格的校准和筛选，我们保留了 158 个验证基准。领域分布突出了 BHI 的多元评估维度，覆盖了推理（20.0%）和多模态（14.8%）等核心领域以及数学和代码生成。此外，供应商分布反映了生态系统多样性，涵盖闭源（OpenAI、Anthropic）、开源（DeepSeek、阿里巴巴）和领域特定模型（美团、小米）。该精心策划的套件为统计分析和 BHI 健康指标计算建立了高保真基础。")

add_separator(doc)

# ============ APPENDIX F ============
add_heading_bilingual(doc, "F  CRITIC Objective Weighting Method", "F  CRITIC 客观加权方法", level=1)

add_bilingual_para(doc,
    "This section provides the detailed mathematical derivation for the CRITIC method used to synthesize the final BHI from the three primary dimensions: SDisc, SAS, and SImp.",
    "本节提供了用于从三个主要维度（SDisc、SAS 和 SImp）合成最终 BHI 的 CRITIC 方法的详细数学推导。")

add_heading_bilingual(doc, "F.1  Statistical Logic of the CRITIC Method", "F.1  CRITIC 方法的统计逻辑", level=2)

add_bilingual_para(doc,
    "The CRITIC method is an objective weighting approach that evaluates the importance of each dimension by considering both its contrast intensity and the conflicting character between indicators. Unlike simpler methods like the entropy weight method, CRITIC accounts for the correlations between variables, ensuring that the weights reflect the total independent information volume carried by each dimension within the global distribution of the benchmarks.",
    "CRITIC 方法是一种客观加权方法，通过考虑每个维度的对比强度和指标之间的冲突特征来评估其重要性。与熵权法等较简单的方法不同，CRITIC 考虑了变量之间的相关性，确保权重反映每个维度在基准全局分布中携带的独立信息总量。")

add_heading_bilingual(doc, "F.2  Mathematical Derivation Steps", "F.2  数学推导步骤", level=2)

add_bilingual_para(doc,
    "Let M denote the total number of benchmarks and n denote the number of evaluation dimensions (n = 3). 1. Dimensional Normalization: Since the calculation logics for SDisc, SAS, and SImp differ, all input dimensions are first mapped to the [0,1] interval via Min-Max normalization to eliminate the influence of scales and units. 2. Calculation of Contrast Intensity: The contrast intensity of the j-th dimension is represented by its standard deviation σ_j. A higher standard deviation indicates a greater variation in values across different benchmarks, suggesting that the dimension provides more discriminative information. 3. Calculation of Conflicting Character: The conflicting character represents the degree of correlation between indicators. Let r_jk denote the Pearson correlation coefficient between the j-th and k-th dimensions. If a dimension is highly correlated with others, it contains more redundant information, and its independent weight should be adjusted accordingly. The conflict index R_j is defined as: R_j = Σ_{k=1}^{n} (1 − r_jk). 4. Calculation of Information Volume: The comprehensive information volume C_j for the j-th dimension is the product of its contrast intensity and its conflicting character: C_j = σ_j · R_j. 5. Determination of Final Weights: The final objective weights w_j are obtained by normalizing the C_j values: w_j = C_j / Σ_{k=1}^{n} C_k.",
    "设 M 表示基准总数，n 表示评估维度数（n = 3）。1. 维度归一化：由于 SDisc、SAS 和 SImp 的计算逻辑不同，所有输入维度首先通过最小-最大归一化映射到 [0,1] 区间，以消除尺度和单位的影响。2. 对比强度计算：第 j 维的对比强度由其标准差 σ_j 表示。标准差越高表示不同基准之间的值变化越大，说明该维度提供更多的区分信息。3. 冲突特征计算：冲突特征表示指标之间的相关程度。设 r_jk 表示第 j 维和第 k 维之间的 Pearson 相关系数。如果一个维度与其他维度高度相关，则包含更多冗余信息，其独立权重应相应调整。冲突指数 R_j 定义为：R_j = Σ_{k=1}^{n} (1 − r_jk)。4. 信息量计算：第 j 维的综合信息量 C_j 是其对比强度和冲突特征的乘积：C_j = σ_j · R_j。5. 最终权重确定：最终客观权重 w_j 通过归一化 C_j 值获得：w_j = C_j / Σ_{k=1}^{n} C_k。")

add_separator(doc)

# ============ APPENDIX G ============
add_heading_bilingual(doc, "G  Robustness under Score Noise and Model Dropout", "G  分数噪声与模型缺失下的稳健性", level=1)

add_bilingual_para(doc,
    "This appendix evaluates the robustness of BHI under two common sources of instability in large-scale benchmark meta-analysis: score perturbation and incomplete model coverage. The first experiment injects controlled Gaussian noise into benchmark scores, simulating reporting variation, measurement uncertainty, and minor evaluation inconsistencies. The second experiment randomly removes evaluated models, simulating incomplete reporting and partial benchmark coverage. For both experiments, we rerun the full BHI computation pipeline after perturbation, including score normalization, LOBO-based model capability calibration, component construction, and final aggregation. We compare the perturbed rankings with the original baseline ranking using Spearman rank correlation, Kendall rank correlation, top-k overlap, and mean rank shift.",
    "本附录评估了 BHI 在大规模基准元分析中两个常见不稳定来源下的稳健性：分数扰动和不完整模型覆盖。第一个实验向基准分数注入受控高斯噪声，模拟报告变异、测量不确定性和轻微评估不一致。第二个实验随机移除已评估模型，模拟不完整报告和部分基准覆盖。对于两个实验，我们在扰动后重新运行完整的 BHI 计算管线，包括分数归一化、基于 LOBO 的模型能力校准、组件构建和最终聚合。我们使用 Spearman 排名相关性、Kendall 排名相关性、前 k 重叠和平均排名偏移将扰动排名与原始基线排名进行比较。")

add_heading_bilingual(doc, "G.1  Robustness to Gaussian Score Noise", "G.1  对高斯分数噪声的稳健性", level=2)

add_bilingual_para(doc,
    "We inject Gaussian noise into the normalized benchmark scores under different noise levels σ ∈ {0.01,0.02,0.05,0.10,0.15,0.20}. Scores are clipped to the valid interval after perturbation. For each noise level, the complete BHI pipeline is recomputed over repeated trials, and the resulting rankings are compared with the clean baseline. Tables 5 and 6 show that BHI rankings remain stable under moderate score perturbations. Under CRITIC weighting, the Spearman correlation remains above 0.95 when σ ≤ 0.05 and above 0.90 when σ ≤ 0.10. Even under the largest tested perturbation, σ = 0.20, the Spearman correlation remains 0.8575 and the BHI score stability remains 0.9458. This suggests that the final ranking is not dominated by small score-level fluctuations. Compared with equal weighting, CRITIC weighting yields consistently higher Spearman correlation, Kendall correlation, score stability, and top-10 overlap across most noise levels. The gap becomes more visible as the noise level increases. This supports the use of CRITIC as a data-driven aggregation mechanism, because it assigns lower influence to less informative or more redundant components rather than treating all dimensions as equally reliable under perturbation.",
    "我们在不同噪声水平 σ ∈ {0.01,0.02,0.05,0.10,0.15,0.20} 下向归一化基准分数注入高斯噪声。扰动后分数被截断到有效区间。对于每个噪声水平，在重复试验中重新计算完整的 BHI 管线，并将结果排名与清洁基线进行比较。表 5 和表 6 显示 BHI 排名在中等分数扰动下保持稳定。在 CRITIC 加权下，当 σ ≤ 0.05 时 Spearman 相关性保持在 0.95 以上，当 σ ≤ 0.10 时保持在 0.90 以上。即使在最大测试扰动 σ = 0.20 下，Spearman 相关性仍为 0.8575，BHI 分数稳定性为 0.9458。这表明最终排名不被小幅分数级波动所主导。与等权加权相比，CRITIC 加权在大多数噪声水平下产生更高的 Spearman 相关性、Kendall 相关性、分数稳定性和前 10 重叠。差距随噪声水平增加而更加明显。这支持了 CRITIC 作为数据驱动聚合机制的使用，因为它将较低的影响力分配给信息量较少或更冗余的组件，而非在扰动下将所有维度视为同等可靠。")

add_heading_bilingual(doc, "G.2  Robustness to Model Dropout", "G.2  对模型缺失的稳健性", level=2)

add_bilingual_para(doc,
    "We further evaluate robustness under incomplete model coverage by randomly removing a fraction of models from the meta-dataset. The dropout ratio ranges from 5% to 60%. After each removal, all benchmark-level metrics are recomputed from the remaining models. Tables 7 and 8 show that the BHI ranking is robust to substantial model removal. Under CRITIC weighting, the Spearman correlation remains above 0.99 at 10% dropout, above 0.97 up to 30% dropout, and above 0.92 even when 60% of models are removed. The top-10 overlap also degrades gradually, indicating that the highest-ranked benchmarks are not highly dependent on a small subset of model reports. The comparison between CRITIC and equal weighting indicates that both aggregation schemes are stable under model dropout, while CRITIC generally provides slightly stronger top-rank preservation in the low-to-moderate dropout regime.",
    "我们通过从元数据集中随机移除一部分模型来进一步评估不完整模型覆盖下的稳健性。缺失率从 5% 到 60%。每次移除后，所有基准级指标从剩余模型重新计算。表 7 和表 8 显示 BHI 排名对大量模型移除具有稳健性。在 CRITIC 加权下，10% 缺失时 Spearman 相关性保持在 0.99 以上，30% 缺失时保持在 0.97 以上，即使 60% 的模型被移除也保持在 0.92 以上。前 10 重叠也逐渐退化，表明排名最高的基准不高度依赖于模型报告的小子集。CRITIC 与等权加权的比较表明，两种聚合方案在模型缺失下均稳定，而 CRITIC 在低至中等缺失情况下通常提供略强的顶级排名保持能力。")

add_separator(doc)

# ============ APPENDIX H ============
add_heading_bilingual(doc, "H  Sensitivity Analysis of Anti-Saturation Parameters", "H  抗饱和度参数的敏感性分析", level=1)

add_bilingual_para(doc,
    "The Anti-Saturation score SAS combines a static resistance component SSta and a dynamic saturation projection component SDyn. In the main experiments, we use a 30-day projection window and combine the two components as SAS = 0.8 · SSta + 0.2 · SDyn. This appendix evaluates whether the final BHI ranking is sensitive to these two design choices.",
    "抗饱和度分数 SAS 结合了静态阻力组件 SSta 和动态饱和投影组件 SDyn。在主要实验中，我们使用 30 天投影窗口并将两个组件组合为 SAS = 0.8 · SSta + 0.2 · SDyn。本附录评估最终 BHI 排名是否对这两个设计选择敏感。")

add_heading_bilingual(doc, "H.1  Sensitivity to the Dynamic Projection Window", "H.1  对动态投影窗口的敏感性", level=2)

add_bilingual_para(doc,
    "We vary the temporal projection window T from 20 to 40 days and recompute the dynamic saturation score as SDyn(b) = max(0, 1 − (Score_b + T · k)). The default setting corresponds to T = 30 days. Table 9 shows that the BHI ranking is nearly invariant to the temporal projection window. Across all tested values, the Spearman correlation remains above 0.9998, the top-10 and top-20 sets are fully preserved, and the average rank shift remains below 0.33. This indicates that the dynamic component captures only a smooth short-term saturation adjustment and does not introduce instability into the final ranking.",
    "我们将时间投影窗口 T 从 20 变化到 40 天，并重新计算动态饱和度为 SDyn(b) = max(0, 1 − (Score_b + T · k))。默认设置对应 T = 30 天。表 9 显示 BHI 排名对时间投影窗口几乎不变。在所有测试值中，Spearman 相关性保持在 0.9998 以上，前 10 和前 20 集完全保留，平均排名偏移保持在 0.33 以下。这表明动态组件仅捕获平滑的短期饱和调整，不会在最终排名中引入不稳定性。")

add_heading_bilingual(doc, "H.2  Sensitivity to Static-Dynamic Weighting", "H.2  对静态-动态加权的敏感性", level=2)

add_bilingual_para(doc,
    "We next vary the relative contribution of SSta and SDyn in the final anti-saturation score. The default setting assigns 80% weight to the static component and 20% weight to the dynamic component. We compare this setting with alternative static weights from 0.6 to 0.9. As shown in Table 10, the final ranking is also robust to the static-dynamic weighting ratio. All tested configurations yield Spearman correlations above 0.999, and the top-20 benchmarks are fully preserved. The largest observed average rank shift is below one position, indicating that the anti-saturation axis is primarily driven by persistent benchmark headroom rather than by short-term trend extrapolation. Overall, these results support the use of a predominantly static formulation for SAS, while retaining a smaller dynamic term to reflect short-term score inflation. The selected 0.8/0.2 configuration provides a conservative balance between present benchmark difficulty and near-future saturation risk.",
    "接下来我们变化 SSta 和 SDyn 在最终抗饱和度分数中的相对贡献。默认设置为静态组件 80% 权重和动态组件 20% 权重。我们将此设置与 0.6 到 0.9 的替代静态权重进行比较。如表 10 所示，最终排名对静态-动态权重比也具有稳健性。所有测试配置产生 0.999 以上的 Spearman 相关性，前 20 基准完全保留。观察到的最大平均排名偏移低于一个位置，表明抗饱和度轴主要由持久的基准提升空间驱动，而非短期趋势外推。总体而言，这些结果支持对 SAS 使用以静态为主的公式，同时保留较小的动态项以反映短期分数膨胀。选定的 0.8/0.2 配置在当前基准难度和近期饱和风险之间提供了保守的平衡。")

add_separator(doc)

# ============ APPENDIX I ============
add_heading_bilingual(doc, "I  Coverage Bias Analysis", "I  覆盖偏差分析", level=1)

add_bilingual_para(doc,
    "A potential concern for any benchmark-level meta-index is that highly evaluated benchmarks may receive higher scores simply because they have broader reporting coverage. This concern is particularly relevant for BHI because the impact component SImp is designed to measure capability-weighted adoption. We therefore analyze whether the final BHI ranking is dominated by benchmark coverage. For each benchmark, we define coverage as the number of models with a valid reported score on that benchmark. We then compute Pearson and Spearman correlations between coverage and each BHI component, as well as the final BHI score.",
    "对于任何基准级元指数，一个潜在的担忧是高度评估的基准可能仅因更广泛的报告覆盖度而获得更高分数。这一担忧对 BHI 尤为相关，因为影响力组件 SImp 旨在衡量能力加权采用度。因此我们分析最终 BHI 排名是否被基准覆盖度主导。对于每个基准，我们将覆盖度定义为在该基准上具有有效报告分数的模型数量。然后我们计算覆盖度与每个 BHI 组件以及最终 BHI 分数之间的 Pearson 和 Spearman 相关性。")

add_bilingual_para(doc,
    "Table 11 shows that coverage is strongly correlated with SImp, which is expected by construction because SImp measures capability-weighted adoption. However, coverage has only weak correlation with SDisc and almost no correlation with SAS. This indicates that broader reporting does not automatically imply stronger discriminative power or greater resistance to saturation. The final BHI score exhibits only moderate correlation with coverage, with Pearson correlation 0.4737 and Spearman correlation 0.4000. This suggests that benchmark coverage contributes to the final score through the impact dimension, but does not dominate the overall ranking. In particular, a benchmark must still provide discriminative resolution and saturation resistance to receive a high BHI score.",
    "表 11 显示覆盖度与 SImp 强相关，这在构造上是预期的，因为 SImp 衡量能力加权采用度。然而，覆盖度与 SDisc 仅弱相关，与 SAS 几乎不相关。这表明更广泛的报告并不自动意味着更强的区分力或更大的抗饱和度。最终 BHI 分数与覆盖度仅呈中等相关，Pearson 相关性 0.4737，Spearman 相关性 0.4000。这表明基准覆盖度通过影响力维度贡献于最终分数，但不主导整体排名。特别是，基准仍必须提供区分分辨率和抗饱和度才能获得高 BHI 分数。")

add_bilingual_para(doc,
    "Figure 15 provides a complementary visualization. Although the fitted trend line is positive, the scatter distribution is broad: benchmarks with similar coverage can receive substantially different BHI scores, and several low-coverage benchmarks obtain competitive BHI values when they exhibit strong discrimination or anti-saturation properties. Conversely, some widely reported benchmarks receive relatively low BHI scores due to saturation or reduced discriminative capacity. These results confirm that BHI is not a proxy for benchmark popularity or reporting frequency. Instead, the final score reflects a balance among technical resolution, remaining headroom, and capability-weighted adoption.",
    "图 15 提供了补充可视化。虽然拟合趋势线为正，但散点分布广泛：覆盖度相似的基准可能获得截然不同的 BHI 分数，而几个低覆盖度基准在表现出强区分度或抗饱和度特性时获得了有竞争力的 BHI 值。反之，一些广泛报告的基准因饱和或区分能力降低而获得相对较低的 BHI 分数。这些结果确认 BHI 不是基准流行度或报告频率的代理。相反，最终分数反映了技术分辨率、剩余提升空间和能力加权采用度之间的平衡。")

add_separator(doc)

# ============ APPENDIX J ============
add_heading_bilingual(doc, "J  Expanded Case Study", "J  扩展案例研究", level=1)

add_bilingual_para(doc,
    "To illustrate how BHI supports benchmark-level diagnosis rather than merely producing a scalar ranking, we analyze three representative cases with distinct metric profiles: Humanity's Last Exam as the strongest overall frontier benchmark in the Knowledge QA category, SimpleQA as a high-resolution short-form factuality benchmark, and ZeroBench as a reasoning benchmark with limited current adoption but substantial remaining headroom. These cases show that the three BHI dimensions capture different types of benchmark value and potential limitations in the current evaluation ecosystem.",
    "为说明 BHI 如何支持基准级诊断而非仅产生标量排名，我们分析了三个具有不同指标特征的代表性案例：Humanity's Last Exam 作为知识问答类别中最强的整体前沿基准，SimpleQA 作为高分辨率短式事实性基准，以及 ZeroBench 作为当前采用有限但剩余提升空间充足的推理基准。这些案例表明三个 BHI 维度捕获了当前评估生态系统中不同类型的基准价值和潜在局限性。")

add_bilingual_para(doc,
    "Humanity's Last Exam. Humanity's Last Exam (HLE) ranks first in our benchmark audit with a BHI score of 0.6846. Its advantage does not come from an extreme value on a single axis, but from a balanced and consistently strong profile across discrimination, saturation resistance, and impact. Specifically, HLE achieves SDisc = 0.6608, SSat = 0.6299, and SImp = 0.7710, indicating that it can separate current models effectively, preserve sufficient remaining challenge space, and maintain broad ecosystem recognition. This makes HLE a suitable anchor for tracking frontier progress in knowledge-intensive and reasoning-oriented evaluation: it combines technical difficulty with stable and influential reporting value in the broader evaluation ecosystem.",
    "Humanity's Last Exam。Humanity's Last Exam（HLE）在我们的基准审计中排名第一，BHI 分数为 0.6846。其优势并非来自单一轴上的极端值，而是来自区分度、抗饱和度和影响力之间均衡且持续强劲的特征。具体而言，HLE 达到 SDisc = 0.6608、SSat = 0.6299 和 SImp = 0.7710，表明它能有效分离当前模型、保留充足的剩余挑战空间，并维持广泛的生态系统认可度。这使 HLE 成为追踪知识密集型和推理导向评估前沿进展的合适锚点：它将技术难度与更广泛评估生态系统中稳定且有影响力的报告价值相结合。")

add_bilingual_para(doc,
    "SimpleQA. SimpleQA obtains a BHI score of 0.5857, reflecting the distinctive role of short-form factuality evaluation in the current benchmark landscape. Its strongest advantage lies in Capability Discrimination, where it reaches SDisc = 0.8130, the highest among the three cases. This suggests that SimpleQA's concise and unambiguous question design reduces interference from complex linguistic formulation and multi-step reasoning, thereby isolating differences in factual answering ability more effectively. At the same time, its saturation score remains moderate (SSat = 0.5043), indicating that it still preserves useful challenge space, though its long-term frontier-tracking capacity is more limited than that of HLE or ZeroBench. Its impact score (SImp = 0.4435) further suggests that SimpleQA has gained meaningful adoption, but has not yet reached the ecosystem coverage of broader frontier benchmarks such as HLE. Therefore, SimpleQA is best viewed as a high-resolution diagnostic tool for factuality evaluation rather than as a standalone benchmark for long-term frontier capability tracking.",
    "SimpleQA。SimpleQA 获得 BHI 分数 0.5857，反映了短式事实性评估在当前基准格局中的独特角色。其最强优势在于能力区分度，达到 SDisc = 0.8130，为三个案例中最高。这表明 SimpleQA 简洁明确的问题设计减少了复杂语言表述和多步推理的干扰，从而更有效地隔离事实回答能力的差异。同时，其饱和度分数保持中等（SSat = 0.5043），表明它仍保留了有用的挑战空间，尽管其长期前沿追踪能力比 HLE 或 ZeroBench 更有限。其影响力分数（SImp = 0.4435）进一步表明 SimpleQA 已获得有意义的采用，但尚未达到 HLE 等更广泛前沿基准的生态系统覆盖度。因此，SimpleQA 最好被视为事实性评估的高分辨率诊断工具，而非长期前沿能力追踪的独立基准。")

add_bilingual_para(doc,
    "ZeroBench. ZeroBench exhibits a different metric structure. Although its impact score is limited (SImp = 0.2498), BHI assigns it a relatively high overall score of 0.5593 due to its strong structural health. Its saturation score reaches SSat = 0.8687, the highest among the three cases, indicating that current models remain far from exhausting its challenge space. Meanwhile, its discrimination score (SDisc = 0.5205) suggests that the benchmark has not collapsed into uniform failure, but can still produce meaningful performance variation across models. This profile identifies ZeroBench as a high-headroom reasoning benchmark whose current adoption underestimates its technical value. For benchmark governance, such cases are important because they reveal high-quality evaluators that deserve more systematic reporting despite relatively weak community momentum.",
    "ZeroBench。ZeroBench 展示了不同的指标结构。虽然其影响力分数有限（SImp = 0.2498），BHI 因其强结构健康度而给予其相对较高的总分 0.5593。其饱和度分数达到 SSat = 0.8687，为三个案例中最高，表明当前模型远未耗尽其挑战空间。同时，其区分度分数（SDisc = 0.5205）表明基准未崩溃为均匀失败，仍能在模型间产生有意义的性能变化。这一特征将 ZeroBench 识别为高提升空间推理基准，其当前采用度低估了其技术价值。对于基准治理，此类案例很重要，因为它们揭示了尽管社区动力相对较弱但值得更系统报告的高质量评估器。")

add_heading_bilingual(doc, "J.1  Domain-Specific Analysis", "J.1  领域特定分析", level=2)

add_bilingual_para(doc,
    "This section provides a domain analysis of 158 benchmarks, offering strategic insights for the design of next-generation evaluation suites and the selection of performance metrics.",
    "本节提供了 158 个基准的领域分析，为下一代评估套件的设计和性能指标的选择提供战略洞察。")

add_heading_bilingual(doc, "J.1.1  Mathematics", "J.1.1  数学", level=3)

add_bilingual_para(doc,
    "In the mathematical domain, a moderate negative correlation (r = −0.22) between Impact and Anti-Saturation suggests that high-impact assets tend toward technical obsolescence. The inherent logical consistency of mathematical reasoning likely facilitates rapid pattern internalization by successor models once these assets are extensively benchmarked. GSM8K: BHI metrics characterize GSM8K as a saturated legacy benchmark. Its minimal anti-saturation score (SAS = 0.0865) and low overall BHI (0.2199) quantify a terminal loss of diagnostic utility for frontier model architectures. Although SDisc remains at 0.3920, the data indicates that GSM8K has transitioned from a research probe into a historical baseline representing past consensus rather than modern technical challenges. FrontierMath: FrontierMath exhibits high technical resilience, defined by a robust anti-saturation score (SAS = 0.6102). While its current impact score (SImp = 0.3339) reflects an early adoption phase, the BHI of 0.5093 highlights its capacity to resist pattern matching. The indicators suggest that by requiring precise mathematical objects, the benchmark maintains the necessary evaluative headroom for future model scaling.",
    "在数学领域，影响力与抗饱和度之间存在中等负相关（r = −0.22），表明高影响力资产倾向于技术过时。数学推理的内在逻辑一致性可能促进后继模型在广泛基准化后快速内化模式。GSM8K：BHI 指标将 GSM8K 表征为饱和的遗留基准。其最低的抗饱和度分数（SAS = 0.0865）和低 BHI 总分（0.2199）量化了前沿模型架构诊断效用的终结性丧失。虽然 SDisc 保持在 0.3920，但数据表明 GSM8K 已从研究探针转变为代表过去共识而非现代技术挑战的历史基线。FrontierMath：FrontierMath 展示了高技术韧性，由稳健的抗饱和度分数（SAS = 0.6102）定义。虽然其当前影响力分数（SImp = 0.3339）反映了早期采用阶段，但 0.5093 的 BHI 突出了其抗模式匹配能力。指标表明，通过要求精确的数学对象，该基准维持了未来模型扩展所需的评估提升空间。")

add_heading_bilingual(doc, "J.1.2  Agent", "J.1.2  智能体", level=3)

add_bilingual_para(doc,
    "The agentic and tool-use domain exhibits a healthy positive correlation (r = 0.30) between Impact and Anti-Saturation, suggesting that the community consensus is anchored towards benchmarks that maintain high technical difficulty and environmental complexity. Tau-bench Series: The Tau-bench series illustrates how metric profiles reveal the divergent evaluation longevity of sub-scenarios. The Airline subset maintains an SAS of 0.4798, indicating robust resistance to templated learning through its multi-constraint game-theoretic designs. In contrast, the Telecom subset, despite its exceptionally high discrimination score (SDisc = 0.9005), exhibits a rapid decay in SAS to 0.1108. This sharp contrast suggests that while its logic is highly precise for ranking, procedural workflows are more susceptible to standardization once model architectures master the underlying patterns. BrowseComp: BrowseComp yields an SDisc of 0.8042 and a BHI of 0.5845, validating the efficacy of its reverse-design methodology. The SAS of 0.3281 shows that constructing queries from rare facts effectively prevents models from relying on internal knowledge or direct keyword matching. By forcing models toward multi-step, creative search strategies in dynamic web environments, the benchmark maintains strong evaluative vitality.",
    "智能体和工具使用领域的影响力与抗饱和度之间呈现健康的正相关（r = 0.30），表明社区共识锚定于维持高技术难度和环境复杂性的基准。Tau-bench 系列：Tau-bench 系列说明了指标特征如何揭示子场景的分化评估寿命。航空子集保持 0.4798 的 SAS，表明通过多约束博弈论设计对模板化学习具有稳健抵抗力。相比之下，电信子集尽管有极高的区分度分数（SDisc = 0.9005），但 SAS 急剧衰减至 0.1108。这种鲜明对比表明，虽然其逻辑在排名方面高度精确，但程序化工作流在模型架构掌握底层模式后更容易被标准化。BrowseComp：BrowseComp 产生 0.8042 的 SDisc 和 0.5845 的 BHI，验证了其逆向设计方法的有效性。0.3281 的 SAS 表明从稀有事实构建查询有效防止模型依赖内部知识或直接关键词匹配。通过迫使模型在动态 Web 环境中采用多步骤创造性搜索策略，该基准维持了强评估活力。")

add_heading_bilingual(doc, "J.1.3  Code", "J.1.3  代码", level=3)

add_bilingual_para(doc,
    "The paradigm of Code evaluation is shifting, evidenced by the decoupling of SImp and SAS, moving from algorithmic puzzles to comprehensive engineering productivity. HumanEval and MBPP: BHI diagnostics confirm the diminishing utility of traditional benchmarks. Low BHI scores (0.2787 for HumanEval; 0.2371 for MBPP) and stagnant SAS (0.1669 and 0.1744, respectively) suggest their focus on function-level algorithms has been neutralized by training data saturation. These metrics indicate that such benchmarks can no longer effectively distinguish genuine reasoning from memorized patterns. SWE-bench: SWE-bench evaluates engineering proficiency through real-world issue resolution in large-scale repositories. Its diagnostic profile, characterized by BHI = 0.5293 and SDisc = 0.6093, stems from an execution-based scoring mechanism. Significant SImp (0.7336) and SAS (0.2754) scores quantify the necessity for a full engineering loop, including fault localization and unit testing. These indicators confirm the benchmark's capacity to target high-level productivity and complex environment navigation.",
    "代码评估的范式正在转变，SImp 与 SAS 的解耦证明了这一点，从算法谜题转向综合工程生产力。HumanEval 和 MBPP：BHI 诊断确认了传统基准的递减效用。低 BHI 分数（HumanEval 0.2787；MBPP 0.2371）和停滞的 SAS（分别为 0.1669 和 0.1744）表明其聚焦于函数级算法已被训练数据饱和所中和。这些指标表明此类基准不再能有效区分真正推理与记忆模式。SWE-bench：SWE-bench 通过大规模仓库中的真实问题解决来评估工程熟练度。其诊断特征为 BHI = 0.5293 和 SDisc = 0.6093，源自基于执行的评分机制。显著的 SImp（0.7336）和 SAS（0.2754）分数量化了完整工程循环（包括故障定位和单元测试）的必要性。这些指标确认了基准针对高级生产力和复杂环境导航的能力。")

add_heading_bilingual(doc, "J.1.4  Subject", "J.1.4  学科", level=3)

add_bilingual_para(doc,
    "The Subject domain is the most susceptible to data contamination, with a median SAS approximately 53.75% lower than the overall median across all evaluation sets. C-Eval: C-Eval faces a similar obsolescence crisis (BHI 0.2098), driven by a critical collapse in its anti-saturation metric (SAS to 0.0935). This indicates significant training data leakage. Consequently, C-Eval has lost its stratifying power (SDisc of 0.2674), evolving into a memorization check rather than a gauge of generalized subject mastery. MMLU-Pro: In contrast, MMLU-Pro (BHI 0.3919) demonstrates the potential for rehabilitating subject-based benchmarks. By increasing reasoning complexity and expanding the option space, it restores discriminatory utility (SDisc of 0.4572), significantly exceeding the domain average. Although its SAS (0.1727) suggests contamination remains a challenge, the recovery in discrimination confirms MMLU-Pro's efficacy in distinguishing capabilities among top-tier models.",
    "学科领域最易受数据污染影响，中位 SAS 比所有评估集的整体中位数低约 53.75%。C-Eval：C-Eval 面临类似的过时危机（BHI 0.2098），由抗饱和度指标的关键崩溃（SAS 降至 0.0935）驱动。这表明存在显著的训练数据泄漏。因此，C-Eval 已丧失其分层能力（SDisc 为 0.2674），演变为记忆检查而非泛化学科掌握的衡量标准。MMLU-Pro：相比之下，MMLU-Pro（BHI 0.3919）展示了恢复学科基准的潜力。通过增加推理复杂性和扩展选项空间，它恢复了区分别用性（SDisc 为 0.4572），显著超过领域平均水平。虽然其 SAS（0.1727）表明污染仍是挑战，但区分度的恢复确认了 MMLU-Pro 在区分顶级模型能力方面的有效性。")

add_heading_bilingual(doc, "J.1.5  Application", "J.1.5  应用", level=3)

add_bilingual_para(doc,
    "The Application domain benchmarks evaluate models under real-world deployment conditions, spanning agent-oriented customer service pipelines, financial information retrieval, and high-stakes medical consultation. The domain exhibits a moderate positive correlation (r = 0.21) between Impact and Anti-Saturation, suggesting that practical deployment pressures tend to favor benchmarks whose task complexity resists rapid memorization. GDPval-AA Elo achieves a competitive BHI of 0.4225 with a well-rounded diagnostic profile. HealthBench Hard stands out within the Application domain for its exceptional anti-saturation properties, recording an SAS of 0.5924 and a BHI of 0.4546. The base HealthBench benchmark presents a notably weaker profile, with a BHI of 0.3524 and an SDisc of 0.4479, indicating insufficient stratifying power at the current performance frontier. These indicators suggest that for frontier-model evaluation purposes, HealthBench Hard has largely displaced the base version as the more informative diagnostic tool.",
    "应用领域基准在真实部署条件下评估模型，涵盖面向智能体的客户服务管线、金融信息检索和高风险医疗咨询。该领域的影响力与抗饱和度之间呈现中等正相关（r = 0.21），表明实际部署压力倾向于支持任务复杂性抵抗快速记忆的基准。GDPval-AA Elo 以 0.4225 的竞争性 BHI 和全面的诊断特征获得认可。HealthBench Hard 因其卓越的抗饱和度特性在应用领域中脱颖而出，SAS 为 0.5924，BHI 为 0.4546。基础 HealthBench 基准呈现明显较弱的特征，BHI 为 0.3524，SDisc 为 0.4479，表明在当前性能前沿分层能力不足。这些指标表明，对于前沿模型评估目的，HealthBench Hard 已在很大程度上取代基础版本成为更具信息量的诊断工具。")

add_separator(doc)

# ============ APPENDIX K ============
add_heading_bilingual(doc, "K  Compute Resources", "K  计算资源", level=1)

add_bilingual_para(doc,
    "All experiments in this work are lightweight post-hoc meta-analytic computations over publicly reported benchmark scores. We do not train, fine-tune, or run inference with any large language model, and no GPU resources are required. The BHI computation, ablation study, robustness analysis, and sensitivity analysis can be reproduced on a standard CPU workstation or laptop with ordinary memory requirements. The total runtime is dominated by data parsing and repeated recomputation of the BHI pipeline rather than model training or benchmark execution.",
    "本工作中的所有实验都是基于公开报告的基准分数的轻量级事后元分析计算。我们不训练、微调或运行任何大语言模型的推理，也不需要 GPU 资源。BHI 计算、消融研究、稳健性分析和敏感性分析可以在标准 CPU 工作站或笔记本电脑上复现，内存需求普通。总运行时间主要由数据解析和 BHI 管线的重复重新计算决定，而非模型训练或基准执行。")

add_separator(doc)

# ============ NOTE ============
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_note.add_run("[References and NeurIPS Paper Checklist are available in the original English document]")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

p_note_cn = doc.add_paragraph()
p_note_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_note_cn.add_run("[参考文献和 NeurIPS 论文清单请参见原英文文档]")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0, 0, 128)
run.font.name = 'SimSun'

# ============ SAVE ============
output_path = '/Users/cishan/Public/Drop Box/基准健康指数看板/BHI_NeurIPS_2026_中英对照.docx'
doc.save(output_path)
print(f'Bilingual document saved to: {output_path}')